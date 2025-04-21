import os
import time
from typing import Dict, List, Union, Optional, Tuple
from docx import Document
import fitz 
from pdf2docx import Converter
from docx2pdf import convert
from publication_chunker import PublicationChunker
from llm_manager import LLMManager
from performance_metrics import PerformanceMetrics

class PublicationTranslator:
    """
    A class to translate publications from any language to English or Arabic.
    """
        
    def __init__(self):
        """
        Initialize the PublicationTranslator using NLLB model.
        """
        self.llm_manager = LLMManager(model_name="facebook/nllb-200-distilled-600M", model_type="seq2seq")
        self.supported_target_languages = ["english", "arabic"]
        self.language_code_map = {
            "english": "eng_Latn",
            "arabic": "arb_Arab"
        }
        self.performance_metrics = PerformanceMetrics()
        self.chunker = PublicationChunker()

    def translate_document(self, file_path: str, target_language: str = "english") -> str:
        """
        Translate a document to the target language while preserving structure.
        
        Args:
            file_path: Path to the document file
            target_language: Target language for translation (english or arabic)
            
        Returns:
            Translated content as string
        """
        target_language = target_language.lower()
        if target_language not in self.supported_target_languages:
            raise ValueError(f"Unsupported target language: {target_language}. Supported languages: {', '.join(self.supported_target_languages)}")
        
        # Get target language code for NLLB
        target_lang_code = self.language_code_map[target_language]
        
        # Check file extension to determine processing method
        file_ext = os.path.splitext(file_path)[1].lower()
        
        if file_ext == '.pdf':
            # Convert PDF to DOCX first, then translate, then convert back to PDF
            return self._translate_pdf_via_docx(file_path, target_lang_code)
        elif file_ext == '.docx':
            return self._translate_docx(file_path, target_lang_code)
        else:
            # For text files or other formats, use the original chunking method
            chunks = self.chunker.chunk_text(file_path, max_tokens=256)  
            translated_chunks = []
            
            # Reset performance metrics for this document
            self.performance_metrics.reset()
            self.performance_metrics.start_tracking()

            for j, chunk in enumerate(chunks):
                print(f"Translating chunk {j+1}/{len(chunks)}...")
                translated_chunk = self._translate_text(chunk, target_lang_code)
                translated_chunks.append(translated_chunk)

            # Stop tracking without assigning the return value
            self.performance_metrics.stop_tracking()
            self.performance_metrics.print_summary()
            
            translated_pages = "\n\n".join(translated_chunks)
            
            return translated_pages
    
    def _translate_docx(self, file_path: str, target_lang_code: str) -> str:
        """
        Translate a docx document while preserving its structure.
        
        Args:
            file_path: Path to the docx file
            target_lang_code: Target language code for NLLB
            
        Returns:
            Translated content as string (for compatibility with the original method)
        """
        try:
            doc = Document(file_path)
            
            # Reset performance metrics for this document
            self.performance_metrics.reset()
            self.performance_metrics.start_tracking()
            
            # Process paragraphs
            total_elements = len(doc.paragraphs) + sum(len(table.rows) for table in doc.tables)
            processed_elements = 0
            
            # Translate paragraphs
            for i, para in enumerate(doc.paragraphs):
                if para.text.strip():  # Skip empty paragraphs
                    processed_elements += 1
                    para.text = self._translate_text(para.text, target_lang_code)
            
            # Translate tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if para.text.strip():  # Skip empty paragraphs
                                processed_elements += 1                             
                                para.text = self._translate_text(para.text, target_lang_code)
            
            # Stop tracking
            self.performance_metrics.stop_tracking()
            self.performance_metrics.print_summary()
            
            # For compatibility with the original method, return a string representation
            # This is only used for the return value, the actual document is preserved in memory
            result = "\n\n".join(p.text for p in doc.paragraphs if p.text.strip())
            
            # Store the translated document for saving
            self._translated_doc = doc
            
            return result
            
        except Exception as e:
            print(f"Error translating docx: {str(e)}")
            return f"[Translation Error: {str(e)}]"
    
    def _translate_text(self, text: str, target_lang_code: str) -> str:
        """
        Translate text using NLLB model while preserving structure and formatting.
        
        Args:
            text: Text to translate
            target_lang_code: Target language code for NLLB
            
        Returns:
            Translated text.
        """
        try:
            start_time = time.time()
               
            response = self.llm_manager.generate_translation(
                text,
                target_lang_code,
                max_length=256,
            )
            
            elapsed_time = time.time() - start_time
            token_count = self.chunker.get_token_count(response)
            
            self.performance_metrics.add_chunk_metrics(token_count, elapsed_time)

            return response

        except Exception as e:
            print(f"Translation error: {str(e)}")
            return f"[Translation Error: {str(e)}]"

    
    def _translate_pdf_via_docx(self, pdf_path: str, target_lang_code: str) -> str:
        """
        Translate a PDF document by converting to DOCX, translating, and converting back.
        
        Args:
            pdf_path: Path to the PDF file
            target_lang_code: Target language code for NLLB
            
        Returns:
            Translated content as string
        """
        try:
            # Create temporary DOCX file
            temp_docx = os.path.splitext(pdf_path)[0] + '_temp.docx'
            
            # Convert PDF to DOCX
            print("Converting PDF to DOCX...")
            cv = Converter(pdf_path)
            cv.convert(temp_docx)
            cv.close()
            
            # Translate the DOCX
            print("Translating DOCX content...")
            translated_content = self._translate_docx(temp_docx, target_lang_code)
            
            # Save translated content to a new DOCX
            translated_docx = os.path.splitext(pdf_path)[0] + '_translated.docx'
            self._translated_doc.save(translated_docx)
            
            # Convert back to PDF
            print("Converting translated DOCX back to PDF...")
            output_pdf = os.path.splitext(pdf_path)[0] + '_translated.pdf'
            convert(translated_docx, output_pdf)
            
            # Clean up temporary files
            if os.path.exists(temp_docx):
                os.remove(temp_docx)
            if os.path.exists(translated_docx):
                os.remove(translated_docx)
            if os.path.exists(output_pdf):
                os.remove(output_pdf)

            return translated_content
            
        except Exception as e:
            print(f"Error translating PDF: {str(e)}")
            return f"[Translation Error: {str(e)}]"

    def save_translated_document(self, translated_content: str, output_path: str) -> str:
        """
        Save the translated content to a file.
        
        Args:
            translated_content: Translated content as string
            output_path: Path to save the translated document
            
        Returns:
            Path to the saved file
        """
        # Get the original file extension
        original_ext = os.path.splitext(output_path)[1].lower()
        
        # If no extension or not supported, default to docx
        if not original_ext or original_ext not in ['.docx', '.pdf']:
            output_path = f"{os.path.splitext(output_path)[0]}.docx"
            original_ext = '.docx'
        
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        
        if original_ext == '.pdf' and hasattr(self, '_translated_doc'):
            pdf_path = f"{os.path.splitext(output_path)[0]}_translated.pdf"
            if os.path.exists(pdf_path):
                return pdf_path
            else:
                # Convert DOCX to PDF if needed
                from docx2pdf import convert
                docx_path = f"{os.path.splitext(output_path)[0]}.docx"
                self._translated_doc.save(docx_path)
                convert(docx_path, output_path)
                if os.path.exists(docx_path):
                    os.remove(docx_path)
                return output_path
        
            
        elif original_ext == '.docx' and hasattr(self, '_translated_doc'):
            # Save the translated DOCX
            self._translated_doc.save(output_path)
            return output_path
            
        else:
            # For text or other formats, create a new DOCX
            doc = Document()
            
            # Split by double newlines to preserve paragraph structure
            paragraphs = translated_content.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    doc.add_paragraph(para_text)
            
            # Save as DOCX
            docx_path = f"{os.path.splitext(output_path)[0]}.docx"
            doc.save(docx_path)
            return docx_path
        